# -*- coding: utf-8 -*-
"""Rebuild the detailed design document in the teacher template structure."""

from __future__ import annotations

from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


PROJECT = Path(__file__).resolve().parents[1]
DOCS = PROJECT / "docs"
SCREENSHOTS = PROJECT / "screenshots" / "app"
ASSETS = DOCS / "design-assets"

OUT_MD = DOCS / "SmartReadAgent_软件详细设计说明书_初稿.md"
OUT_DOCX = DOCS / "SmartReadAgent_软件详细设计说明书_初稿.docx"
ALT_DOCX = DOCS / "SmartReadAgent_软件详细设计说明书_模板版.docx"


FONT_PATHS = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simsun.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    for path in FONT_PATHS:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(text: str, chars: int) -> str:
    lines: list[str] = []
    for raw in text.split("\n"):
        if len(raw) <= chars:
            lines.append(raw)
        else:
            lines.extend(textwrap.wrap(raw, width=chars))
    return "\n".join(lines)


def draw_centered(draw: ImageDraw.ImageDraw, xy, text: str, fill, fnt, anchor="mm"):
    draw.text(xy, text, fill=fill, font=fnt, anchor=anchor, align="center", spacing=6)


def rounded_box(draw, xy, fill, outline, radius=22, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill=(45, 55, 72), width=4):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - sign * 16, y2 - 9), (x2 - sign * 16, y2 + 9)]
    else:
        sign = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - sign * 16), (x2 + 9, y2 - sign * 16)]
    draw.polygon(points, fill=fill)


def make_use_case_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    label_font = font(28)
    small_font = font(23)
    draw_centered(draw, (750, 58), "SmartRead Agent 顶层用例图", "#111827", title_font)

    # Actors
    for x, name in [(150, "学生用户"), (1350, "小组演示者")]:
        draw.ellipse((x - 38, 140, x + 38, 216), outline="#1d4ed8", width=4)
        draw.line((x, 216, x, 330), fill="#1d4ed8", width=4)
        draw.line((x - 70, 260, x + 70, 260), fill="#1d4ed8", width=4)
        draw.line((x, 330, x - 65, 440), fill="#1d4ed8", width=4)
        draw.line((x, 330, x + 65, 440), fill="#1d4ed8", width=4)
        draw_centered(draw, (x, 490), name, "#111827", label_font)

    use_cases = [
        (445, 170, "输入/选择\n示例文本"),
        (750, 170, "生成摘要\n关键词"),
        (1055, 170, "LiteRT 端侧\n句子分析"),
        (445, 390, "Agent 问答"),
        (750, 390, "知识卡片\n复习题"),
        (1055, 390, "历史记录\n恢复/清空"),
        (750, 625, "现场演示\n功能流程"),
    ]
    for x, y, text in use_cases:
        draw.ellipse((x - 155, y - 72, x + 155, y + 72), fill="#dbeafe", outline="#2563eb", width=3)
        draw_centered(draw, (x, y), text, "#1e3a8a", label_font)

    for target in [(290, 170), (290, 390), (595, 625)]:
        arrow(draw, (220, 300), target)
    for target in [(905, 170), (905, 390), (905, 625)]:
        arrow(draw, (1280, 300), target)
    arrow(draw, (600, 170), (595, 170))
    arrow(draw, (905, 170), (900, 170))
    arrow(draw, (600, 390), (595, 390))
    arrow(draw, (905, 390), (900, 390))
    draw.text((90, 820), "说明：本项目当前 Agent 为本地规则型阅读 Agent，不依赖外部大模型接口；LiteRT 模型在 Android 端侧完成句子重要性评分。", fill="#374151", font=small_font)
    img.save(path)


def make_object_model_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 920), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    label_font = font(25)
    small_font = font(21)
    draw_centered(draw, (750, 55), "SmartRead Agent 数据对象模型", "#111827", title_font)

    boxes = {
        "ArticleAnalysis": (540, 125, 960, 270, "ArticleAnalysis\n原文、摘要、关键词、分点摘要\n句子数、字符数、LiteRT状态"),
        "SentenceImportance": (80, 360, 450, 505, "SentenceImportance\nsentence / score / level\nsource / index"),
        "ChatMessage": (565, 360, 935, 505, "ChatMessage\nrole / content\nUSER 或 AGENT"),
        "KnowledgeCard": (1050, 360, 1420, 505, "KnowledgeCard\ntitle / type / content"),
        "QuizQuestion": (315, 625, 685, 770, "QuizQuestion\nquestion\nreferenceAnswer"),
        "HistoryRecord": (820, 625, 1190, 770, "HistoryRecord\nid / title / preview / savedAt\noriginalText / summary / keywords"),
    }
    for name, (x1, y1, x2, y2, text) in boxes.items():
        rounded_box(draw, (x1, y1, x2, y2), "#e0f2fe", "#0284c7")
        draw_centered(draw, ((x1 + x2) // 2, (y1 + y2) // 2), text, "#0f172a", label_font)

    arrow(draw, (540, 270), (450, 360))
    arrow(draw, (700, 270), (700, 360))
    arrow(draw, (960, 270), (1050, 360))
    arrow(draw, (665, 505), (500, 625))
    arrow(draw, (835, 505), (1000, 625))
    draw.text((85, 845), "说明：ArticleAnalysis 是主分析结果对象，其他对象围绕问答、卡片、复习题、历史记录和 LiteRT 分析展开。", fill="#374151", font=small_font)
    img.save(path)


def _make_module_diagram_old(path: Path) -> None:
    img = Image.new("RGB", (1500, 920), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    label_font = font(24)
    small_font = font(20)
    draw_centered(draw, (750, 55), "SmartRead Agent 模块关系图", "#111827", title_font)
    modules = [
        (60, 160, 310, 255, "文本输入模块"),
        (395, 160, 645, 255, "TextAnalyzer\n摘要/关键词"),
        (730, 160, 980, 255, "ArticleAnalysis\n分析结果"),
        (1065, 160, 1315, 255, "Compose UI\n结果展示"),
        (395, 395, 645, 490, "SentenceImportance\nAnalyzer"),
        (730, 395, 980, 490, "LiteRT Classifier\n.tflite 推理"),
        (395, 630, 645, 725, "LocalReadingAgent\n本地规则问答"),
        (730, 630, 980, 725, "HistoryRepository\n历史记录"),
        (1065, 630, 1315, 725, "Knowledge/Quiz\n复习材料"),
    ]
    for x1, y1, x2, y2, text in modules:
        rounded_box(draw, (x1, y1, x2, y2), "#ecfdf5", "#059669")
        draw_centered(draw, ((x1 + x2) // 2, (y1 + y2) // 2), text, "#064e3b", label_font)

    arrow(draw, (310, 207), (395, 207))
    arrow(draw, (645, 207), (730, 207))
    arrow(draw, (980, 207), (1065, 207))
    arrow(draw, (855, 255), (520, 395))
    arrow(draw, (645, 442), (730, 442))
    arrow(draw, (855, 490), (1065, 650))
    arrow(draw, (855, 255), (520, 630))
    arrow(draw, (645, 677), (1065, 677))
    arrow(draw, (855, 255), (855, 630))
    arrow(draw, (980, 677), (1065, 677))
    draw.text((75, 835), "说明：主流程先生成 ArticleAnalysis，再分发给 LiteRT、Agent、知识卡片和历史记录模块；模型失败时有规则兜底。", fill="#374151", font=small_font)
    img.save(path)


def make_module_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 920), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    label_font = font(23)
    small_font = font(20)
    lane_font = font(18)
    line_color = "#334155"

    def draw_arrow_head(start, end, fill=line_color):
        x1, y1 = start
        x2, y2 = end
        if abs(x2 - x1) >= abs(y2 - y1):
            sign = 1 if x2 > x1 else -1
            points = [(x2, y2), (x2 - sign * 16, y2 - 9), (x2 - sign * 16, y2 + 9)]
        else:
            sign = 1 if y2 > y1 else -1
            points = [(x2, y2), (x2 - 9, y2 - sign * 16), (x2 + 9, y2 - sign * 16)]
        draw.polygon(points, fill=fill)

    def connector(points, fill=line_color, width=4):
        draw.line(points, fill=fill, width=width, joint="curve")
        draw_arrow_head(points[-2], points[-1], fill)

    draw_centered(draw, (750, 55), "SmartRead Agent 模块关系图", "#111827", title_font)
    draw.text((70, 120), "主流程", fill="#64748b", font=lane_font)
    draw.text((70, 335), "分析结果分发", fill="#64748b", font=lane_font)
    draw.text((70, 555), "本地能力模块", fill="#64748b", font=lane_font)

    modules = [
        (90, 155, 300, 240, "文本输入模块"),
        (390, 155, 625, 240, "TextAnalyzer\n摘要/关键词"),
        (715, 155, 950, 240, "ArticleAnalysis\n分析结果"),
        (1060, 155, 1295, 240, "Compose UI\n结果展示"),
        (140, 405, 400, 500, "SentenceImportance\nAnalyzer"),
        (620, 405, 880, 500, "LocalReadingAgent\n本地规则问答"),
        (1100, 405, 1360, 500, "Knowledge/Quiz\n复习材料"),
        (140, 635, 400, 730, "LiteRT Classifier\n.tflite 推理"),
        (620, 635, 880, 730, "HistoryRepository\n历史记录"),
    ]
    for x1, y1, x2, y2, text in modules:
        rounded_box(draw, (x1, y1, x2, y2), "#ecfdf5", "#059669")
        draw_centered(draw, ((x1 + x2) // 2, (y1 + y2) // 2), text, "#064e3b", label_font)

    # Main pipeline.
    connector([(300, 198), (390, 198)])
    connector([(625, 198), (715, 198)])
    connector([(950, 198), (1060, 198)])

    # Distribution bus from ArticleAnalysis. The bus keeps all branch lines away
    # from module boxes and avoids the earlier diagonal crossings.
    draw.line([(832, 240), (832, 340), (270, 340), (1230, 340)], fill=line_color, width=4, joint="curve")
    connector([(270, 340), (270, 405)])
    connector([(750, 340), (750, 405)])
    connector([(1230, 340), (1230, 405)])

    # Local processing relationships.
    connector([(270, 500), (270, 635)])
    connector([(750, 500), (750, 635)])
    connector([(880, 452), (1100, 452)])

    draw.text(
        (75, 835),
        "说明：主链路先生成 ArticleAnalysis，再通过分发线交给端侧分析、问答、复习材料和历史记录模块；折线连接只表示模块调用或数据流向。",
        fill="#374151",
        font=small_font,
    )
    img.save(path)


def make_use_case_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    label_font = font(27)
    small_font = font(21)
    body_font = font(24)
    muted = "#475569"

    draw_centered(draw, (750, 58), "SmartRead Agent 顶层用例图", "#111827", title_font)

    def actor_card(x1, y1, x2, y2, title, items, fill):
        rounded_box(draw, (x1, y1, x2, y2), fill, "#2563eb", radius=24, width=3)
        draw_centered(draw, ((x1 + x2) // 2, y1 + 42), title, "#1e3a8a", label_font)
        y = y1 + 92
        for idx, item in enumerate(items, start=1):
            rounded_box(draw, (x1 + 36, y, x2 - 36, y + 62), "#ffffff", "#93c5fd", radius=16, width=2)
            draw.text((x1 + 58, y + 16), f"{idx}. {item}", fill="#0f172a", font=body_font)
            y += 82

    student_items = [
        "输入或选择示例文本",
        "生成摘要、关键词和分点摘要",
        "查看 LiteRT 端侧句子分析",
        "进行 Agent 问答",
        "查看知识卡片和复习题",
        "恢复或清空历史记录",
    ]
    demo_items = [
        "按演示脚本走完整功能流程",
        "展示摘要、端侧分析和问答结果",
        "展示复习材料与历史记录",
    ]

    actor_card(110, 140, 690, 720, "学生用户", student_items, "#dbeafe")
    actor_card(810, 140, 1390, 490, "小组演示者", demo_items, "#e0f2fe")

    rounded_box(draw, (810, 560, 1390, 720), "#ecfdf5", "#059669", radius=24, width=3)
    draw_centered(draw, (1100, 605), "系统边界", "#065f46", label_font)
    draw.text((855, 655), "SmartRead Agent Android App", fill="#0f172a", font=body_font)
    draw.text((855, 690), "输入、分析、复习、历史记录均在移动端完成展示", fill=muted, font=small_font)

    draw.text(
        (90, 820),
        "说明：用例按参与者分组展示，每一项均对应 V1.0 演示中的可操作功能。",
        fill="#374151",
        font=small_font,
    )
    img.save(path)


def make_object_model_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 920), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    label_font = font(25)
    small_font = font(20)
    body_font = font(23)

    draw_centered(draw, (750, 55), "SmartRead Agent 数据对象模型", "#111827", title_font)

    def object_card(x1, y1, x2, y2, title, fields, note, fill="#e0f2fe"):
        rounded_box(draw, (x1, y1, x2, y2), fill, "#0284c7", radius=22, width=3)
        draw_centered(draw, ((x1 + x2) // 2, y1 + 38), title, "#0f172a", label_font)
        y = y1 + 78
        for field in fields:
            draw.text((x1 + 28, y), field, fill="#0f172a", font=body_font)
            y += 34

    object_card(
        420,
        125,
        1080,
        325,
        "ArticleAnalysis  主分析结果",
        ["originalText / oneSentenceSummary / keywords", "bulletSummary / sentenceCount / characterCount", "sentenceImportanceList / createdAt"],
        "由 TextAnalyzer 生成，是后续展示、问答和复习材料的基础。",
        "#dbeafe",
    )
    object_card(
        95,
        370,
        455,
        585,
        "SentenceImportance",
        ["sentence", "score", "level", "source / index"],
        "LiteRT 或规则分析后的句子重要性结果。",
    )
    object_card(
        570,
        370,
        930,
        585,
        "ChatMessage",
        ["role", "content", "timestamp"],
        "Agent 问答界面中的单条对话消息。",
    )
    object_card(
        1045,
        370,
        1405,
        585,
        "KnowledgeCard",
        ["title", "type", "content"],
        "从摘要、关键词和复习重点整理出的卡片。",
    )
    object_card(
        330,
        660,
        690,
        820,
        "QuizQuestion",
        ["question", "referenceAnswer"],
        "根据文章内容生成的复习题。",
        "#ecfdf5",
    )
    object_card(
        810,
        660,
        1230,
        820,
        "HistoryRecord",
        ["id / title / preview / savedAt", "originalText / summary / keywords"],
        "本地历史记录保存和恢复使用的数据结构。",
        "#ecfdf5",
    )

    draw.text(
        (85, 875),
        "说明：对象模型采用分区卡片展示字段和用途；主对象为 ArticleAnalysis，其余对象围绕展示、问答、复习和历史记录展开。",
        fill="#374151",
        font=small_font,
    )
    img.save(path)


def make_diagrams() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths = {
        "use_case": ASSETS / "SmartReadAgent_顶层用例图.png",
        "object": ASSETS / "SmartReadAgent_对象模型图.png",
        "module": ASSETS / "SmartReadAgent_模块关系图.png",
    }
    make_use_case_diagram(paths["use_case"])
    make_object_model_diagram(paths["object"])
    make_module_diagram(paths["module"])
    return paths


def screenshot(name: str) -> Path:
    path = SCREENSHOTS / name
    if not path.exists():
        raise FileNotFoundError(path)
    return path


def design_markdown() -> str:
    return """# SmartRead Agent 系统 软件详细设计说明书

Version 1.0

文档名称：SmartReadAgent_软件详细设计说明书_模板版.docx

## 修订历史记录

| 日期 | 版本号 | 修改说明 | 修改人 | 核准人 |
|---|---|---|---|---|
| 2026-05-11 | V0.1 | 项目立项、目录和 Git/Gitee/GitHub 初始化 | 林立洲 | 林立洲 |
| 2026-05-12 | V0.2 | 完成文本输入、本地摘要、关键词和分点摘要 MVP | 林立洲 | 林立洲 |
| 2026-05-12 | V0.3 | 增加 Agent 问答、知识卡片和复习题 | 林立洲 | 林立洲 |
| 2026-05-12 | V0.4 | 增加 LiteRT 端侧句子重要性分析模型 | 林立洲 | 林立洲 |
| 2026-05-12 | V0.5 | 增加本地历史记录保存、恢复和清空 | 林立洲 | 林立洲 |
| 2026-05-13 | V1.0 | 收口期末演示稳定版、APK、截图、测试和说明材料 | 林立洲、江轩宇 | 林立洲 |
| 2026-06-29 | V1.0 初稿 | 按课程模板整理软件详细设计说明书 | 林立洲 | 林立洲 |

## 目录

1 引言

1.1 目的与范围

1.2 预期的读者

1.3 系统的范围

1.4 参考资料

1.5 术语、缩写词

1.6 模块命名规则

2 建议的系统

2.1 建议系统概述

2.2 功能性需求概述

2.3 非功能性需求

2.3.1 用户界面与人员因素

2.3.2 错误处理与极端情况

2.3.3 质量要求

2.3.4 用例模型

2.3.5 对象模型

3 模块总汇

3.1 模块总汇表

3.2 模块关系图

3.3 模块设计清单

3.4 附录：数据存储结构设计

## 1 引言

### 1.1 目的与范围

本文档是 SmartRead Agent 系统的详细设计说明书，按照课程提供的“软件详细设计说明书模板”结构编写。文档用于说明系统设计目标、功能范围、非功能需求、用例模型、对象模型、模块关系、模块设计清单和数据存储结构，为《软件项目研发实践（3）》期末演示后的电子版与纸质版提交提供依据。

SmartRead Agent 当前版本为 V1.0 期末演示稳定版，已生成课程演示用 debug APK：`release/SmartReadAgent-v1.0-debug.apk`。该 APK 用于课程验收、真机测试和现场演示，不作为应用商店正式商业发布版本。本文档只描述当前已实现的功能，不把 OCR、云端大模型、登录、服务器、多设备同步等后续设想写成已实现内容。

本文档覆盖 Android/Kotlin 客户端、Jetpack Compose 界面、本地文本分析、本地规则型阅读 Agent、LiteRT 端侧句子重要性分析、本地历史记录和演示部署流程。重点突出移动端机器学习在应用中的真实整合方式，即模型文件随 APK 集成并在 Android 端侧完成推理。

### 1.2 预期的读者

本文档的预期读者包括课程指导老师、期末演示评审人员、小组成员、后续查看仓库的同学，以及需要理解系统模块设计和代码结构的人。读者不需要提前阅读所有源代码，但应能通过本文档理解系统主要模块、数据对象和运行流程。对于评分人员，本文档也可作为核对项目是否真实实现 LiteRT、Agent、历史记录和移动端演示能力的依据。

### 1.3 系统的范围

SmartRead Agent 面向学生阅读教材、课堂资料、技术文章和项目说明文档的场景，提供移动端阅读辅助能力。系统当前范围包括：文本输入、示例文本、本地摘要、关键词提取、分点摘要、LiteRT 端侧句子重要性分析、Agent 问答、知识卡片、复习题、本地历史记录保存、历史记录恢复和清空历史记录。

系统当前不包括：用户注册登录、云端同步、服务器后台、OCR 图片识别、语音输入、多设备同步、商业发布版签名包和真实云端大模型 API。后续如果继续完善，可以在当前本地能力稳定的基础上逐步扩展。

### 1.4 参考资料

- 《2023级软工期末考核说明.pdf》。
- 《2023级软工期末课程演示流程.pdf》。
- 《2023级期末演示评分表.xlsx》。
- `README.md`、`docs/V0.4_LiteRT端侧模型集成说明.md`、`docs/关键代码与实现方法说明.md`。
- Android / Kotlin / Jetpack Compose 相关开发资料。
- TensorFlow Lite / LiteRT Android 端侧推理相关资料。

### 1.5 术语、缩写词

| 术语 | 说明 |
|---|---|
| SmartRead Agent | 本项目名称，表示面向阅读场景的智能辅助 App |
| LiteRT / TensorFlow Lite | 移动端轻量模型推理框架 |
| 本地规则型阅读 Agent | 当前版本 Agent 实现方式，主要使用本地规则和文章分析结果 |
| ArticleAnalysis | 文章分析结果数据对象 |
| SentenceImportance | 句子重要性评分结果对象 |
| APK | Android 安装包 |
| SharedPreferences | Android 本地轻量键值存储 |

### 1.6 模块命名规则

系统代码采用 Kotlin 命名规范。数据对象使用名词型 PascalCase，例如 `ArticleAnalysis`、`SentenceImportance`、`KnowledgeCard`、`HistoryRecord`。功能对象或仓库对象使用职责名加后缀，例如 `TextAnalyzer`、`SentenceImportanceAnalyzer`、`SentenceImportanceClassifier`、`LocalReadingAgent`、`HistoryRepository`。Compose 页面函数使用页面含义命名，例如 `AgentChatScreen`、`KnowledgeCardsScreen`、`HistoryScreen`。资源文件使用能够说明功能的名称，例如 `sentence_importance_model.tflite`。

## 2 建议的系统

### 2.1 建议系统概述

SmartRead Agent 是一个 Android/Kotlin 移动端 AI 阅读辅助应用。用户可以在首页输入文本或点击示例文本，然后点击“开始智能分析”。系统会在本地生成一句话总结、关键词和分点摘要，同时调用 LiteRT 模型对句子重要性进行端侧评分。分析完成后，用户可以进入 Agent 问答页面围绕当前文章提问，也可以进入知识卡片页面查看复习卡片和复习题，还可以在历史记录页面恢复最近的分析结果。

系统设计强调“轻量、稳定、可解释、适合课堂演示”。当前版本的 Agent 问答基于当前文章分析结果和本地规则实现，不依赖外部大模型接口，因此在没有网络的情况下也能完成演示。LiteRT 模块是系统移动机器学习部分的核心，模型文件放入 Android assets，运行时由 `SentenceImportanceClassifier` 加载并推理。

### 2.2 功能性需求概述

系统功能性需求如下：

1. 用户可以手动输入待分析文本。
2. 用户可以选择内置示例文本，保证课堂演示稳定。
3. 系统可以生成一句话总结、关键词和分点摘要。
4. 系统可以显示字符数、句子数等基础统计信息。
5. 系统可以调用 LiteRT 端侧模型，对文章句子进行重要性评分。
6. 系统可以在模型加载失败时使用规则评分兜底，避免演示中断。
7. 系统可以进入 Agent 问答页面，支持快捷问题和手动提问。
8. 系统可以生成知识卡片和复习题。
9. 系统可以保存最近 12 条历史分析记录。
10. 用户可以恢复历史记录，也可以清空历史记录。

### 2.3 非功能性需求

系统需要满足现场可演示、离线可用、响应及时、界面清晰、异常可控和材料可追溯等非功能需求。由于课程演示每组不超过 15 分钟，App 操作路径必须简洁，示例文本必须可直接触发完整功能链路。由于教室网络环境不一定稳定，核心功能不依赖网络。

### 2.3.1 用户界面与人员因素

界面采用 Jetpack Compose 实现，整体为移动端纵向滚动布局。首页突出项目名称、版本号、输入区、示例文本按钮和“开始智能分析”按钮。分析结果以卡片方式展示，包含摘要、关键词、LiteRT 评分和功能入口。Agent、知识卡片和历史记录作为独立页面，便于演示时逐个展示。

用户角色主要是学生用户和小组演示者。学生用户关注能否快速理解文章；演示者关注流程是否稳定、老师是否能看清楚功能和机器学习亮点。江轩宇在现场负责手机端实际操作，林立洲负责同步讲解和答辩。

界面截图已直接插入本文档对应位置，不留空白页。首页、摘要结果、LiteRT、Agent、知识卡片和历史记录截图均来自 `screenshots/app/` 目录。

### 2.3.2 错误处理与极端情况

系统对空文本、模型加载失败、历史记录 JSON 解析失败和无分析结果进入子页面等情况进行了处理。空文本分析时不生成结果，而是提示用户先输入内容；LiteRT 模型加载或推理失败时，系统使用规则评分兜底；历史记录解析失败时返回空列表；用户未完成分析就进入 Agent 或知识卡片页面时，页面显示友好提示。

这些处理的目的不是追求复杂异常体系，而是保证课程现场演示不断线。只要 App 能打开，用户就能通过示例文本稳定走完整流程。

### 2.3.3 质量要求

| 主要质量属性 | 详细要求 |
|---|---|
| 正确性 | 摘要、关键词、LiteRT 分析、Agent 回答、知识卡片和历史记录应基于当前输入文本生成，不展示与当前文章无关的内容。 |
| 健壮性 | 空输入、模型加载失败、历史记录为空、历史 JSON 解析失败等情况不应导致 App 崩溃。 |
| 兼容性 | APK 可在 Android 模拟器和安卓真机上安装运行；课程演示用 debug APK 不作为商业发布版。 |

### 2.3.4 用例模型

参与者(Actor)汇总：

| 参与者 | 说明 |
|---|---|
| 学生用户 | 输入或选择文本，查看摘要、LiteRT 分析、Agent 回答、知识卡片和历史记录。 |
| 小组演示者 | 在课堂现场操作 App，按固定流程展示功能和机器学习整合效果。 |
| 老师 | 查看系统演示、提出问题、根据评分表评价应用创新性、功能丰富度和机器学习整合程度。 |

用例汇总：

| 用例名称 | 参与者 | 简要说明 |
|---|---|---|
| 输入或选择文本 | 学生用户、小组演示者 | 用户输入待分析文本或点击示例文本。 |
| 生成阅读分析结果 | 学生用户、小组演示者 | 系统生成摘要、关键词、分点摘要和基础统计。 |
| 执行 LiteRT 端侧分析 | 学生用户、小组演示者 | 系统对句子进行重要性评分并显示来源。 |
| 使用 Agent 问答 | 学生用户、小组演示者 | 用户围绕当前文章点击快捷问题或手动提问。 |
| 查看知识卡片和复习题 | 学生用户、小组演示者 | 系统根据分析结果生成复习材料。 |
| 查看和恢复历史记录 | 学生用户、小组演示者 | 用户查看最近分析并恢复其中一条。 |

顶层用例图：`docs/design-assets/SmartReadAgent_顶层用例图.png`

用例描述：

1. 输入或选择文本：用户在首页文本框输入内容，或点击示例文本按钮。系统更新输入框内容并等待分析。
2. 生成阅读分析结果：用户点击“开始智能分析”，系统调用 `TextAnalyzer.analyze` 生成 `ArticleAnalysis`。
3. 执行 LiteRT 端侧分析：系统提取句子特征并调用 `SentenceImportanceClassifier`，在结果页显示分数和来源。
4. 使用 Agent 问答：用户进入 Agent 页面，选择快捷问题或输入问题，系统调用 `LocalReadingAgent.answerQuestion` 返回回答。
5. 查看知识卡片和复习题：用户进入知识卡片页面，系统根据当前 `ArticleAnalysis` 生成 `KnowledgeCard` 和 `QuizQuestion`。
6. 查看和恢复历史记录：用户进入历史记录页面，点击某条记录后恢复原文并重新分析。

### 2.3.5 对象模型

数据实体对象：

- `ArticleAnalysis`：系统主分析结果对象，保存原文、一句话总结、关键词、分点摘要、句子数、字符数、句子重要性列表和模型状态。
- `SentenceImportance`：LiteRT 或规则兜底输出的句子重要性结果。
- `ChatMessage`：Agent 问答页面中的单条消息。
- `KnowledgeCard`：知识卡片页面中的复习卡片。
- `QuizQuestion`：复习题对象，包含题目和参考答案。
- `HistoryRecord`：历史记录对象，保存恢复分析所需的信息。

数据相关对象结构：`docs/design-assets/SmartReadAgent_对象模型图.png`

## 3 模块总汇

### 3.1 模块总汇表

| 模块名称 | 类名称 / 对象名称 | 主要职责 |
|---|---|---|
| 文本输入模块 | Compose 首页、SampleTexts | 处理手动输入、示例文本选择、清空和开始分析。 |
| 本地摘要与关键词模块 | TextAnalyzer | 切分句子、提取关键词、生成一句话总结和分点摘要。 |
| LiteRT 端侧分析模块 | SentenceImportanceAnalyzer、SentenceImportanceClassifier | 提取 5 维特征，加载 .tflite 模型并输出句子重要性评分。 |
| Agent 问答模块 | LocalReadingAgent | 根据当前文章分析结果进行本地规则型问答。 |
| 知识卡片与复习题模块 | LocalReadingAgent.generateKnowledgeCards / generateQuizQuestions | 生成复习卡片和简答题。 |
| 历史记录模块 | HistoryRepository | 使用 SharedPreferences 保存、读取、恢复和清空历史记录。 |
| UI 展示模块 | SmartReadAgentApp、AgentChatScreen、KnowledgeCardsScreen、HistoryScreen | 组织首页、结果页、Agent 页面、知识卡片页面和历史页面。 |

### 3.2 模块关系图

模块关系图：`docs/design-assets/SmartReadAgent_模块关系图.png`

系统主流程为：文本输入模块收集原文，交给 `TextAnalyzer` 得到 `ArticleAnalysis`；分析结果继续进入 LiteRT 端侧分析模块、Agent 问答模块、知识卡片模块和历史记录模块；最终由 Compose UI 展示给用户。LiteRT 模型异常时，系统使用规则评分兜底，不影响摘要和其他功能。

### 3.3 模块设计清单

#### 模块一：文本输入模块

模块职责：提供文本输入、示例文本选择、清空和开始分析入口。输入为用户手动输入或内置示例文本；输出为传给 `TextAnalyzer` 的原始字符串。异常处理包括空文本提示和清空后状态重置。

核心界面截图：`screenshots/app/V1.0首页_20260512.png`、`screenshots/app/V1.0文本输入_20260512.png`

#### 模块二：本地摘要与关键词模块

模块职责：生成一句话总结、关键词、分点摘要、字符数和句子数。核心对象为 `TextAnalyzer`。它先切分句子，再按关键词命中、句子长度和位置给句子评分。关键词提取结合领域词、英文技术词和中文短语窗口。

核心界面截图：`screenshots/app/V1.0摘要结果_20260512.png`

#### 模块三：LiteRT 端侧句子重要性分析模块

模块职责：体现移动端机器学习整合。核心对象为 `SentenceImportanceAnalyzer` 和 `SentenceImportanceClassifier`。输入为句子、关键词、句子位置和总句数，输出为 `SentenceImportance`。模型输入 5 维特征：句子长度归一化、关键词重合度、位置分数、标点提示分数、总结提示词分数。模型文件为 `android/app/src/main/assets/sentence_importance_model.tflite`。

异常处理：如果模型加载失败或推理失败，调用 `fallbackScore` 生成规则评分，并显示“规则评分兜底”。这保证现场演示不会因为模型异常中断。

核心界面截图：`screenshots/app/V1.0LiteRT端侧分析_20260512.png`

#### 模块四：Agent 问答模块

模块职责：围绕当前文章分析结果回答用户问题。当前版本统一说明为本地规则型阅读 Agent，回答内容主要来自摘要、关键词、复习题和知识卡片等本地分析结果。核心对象为 `LocalReadingAgent`。它根据问题意图识别返回概括、关键词、复习重点、复习题、知识卡片提示或综合回答。

核心界面截图：`screenshots/app/V1.0Agent问答_20260512.png`、`screenshots/app/V1.0Agent快捷问题_20260512.png`

#### 模块五：知识卡片与复习题模块

模块职责：把分析结果转换为复习材料。输入为 `ArticleAnalysis`，输出为 `KnowledgeCard` 和 `QuizQuestion`。知识卡片包括概念、观点和复习类内容；复习题以简答题和参考答案形式展示。

核心界面截图：`screenshots/app/V1.0知识卡片_20260512.png`、`screenshots/app/V1.0复习题_20260512.png`

#### 模块六：历史记录模块

模块职责：保存最近 12 条分析记录，并支持恢复和清空。核心对象为 `HistoryRepository`。存储方式为 SharedPreferences 中的 JSON 数组，不使用数据库。该方案适合当前课程演示规模，后续如扩展为正式产品可迁移至 Room。

核心界面截图：`screenshots/app/V1.0历史记录页_20260512.png`、`screenshots/app/V1.0历史记录恢复_20260512.png`

#### 模块七：UI 展示模块

模块职责：使用 Jetpack Compose 组织首页、结果页、Agent 页面、知识卡片页面和历史记录页面。页面切换由 `Screen` 枚举控制，UI 状态由 Compose 的 `remember` 和 `mutableStateOf` 管理。界面强调课程演示时的可读性和操作路径清晰。

### 3.4 附录：数据存储结构设计

| 列名 | 数据类型 | 是否为空 | 说明 |
|---|---|---|---|
| id | Long | 否 | 历史记录唯一标识，使用保存时间戳。 |
| title | String | 否 | 历史记录标题，默认取一句话总结前 28 个字符。 |
| preview | String | 否 | 原文预览，列表页展示使用。 |
| savedAt | Long | 否 | 保存时间戳。 |
| displayTime | String | 否 | 格式化后的展示时间。 |
| originalText | String | 否 | 原始文本，用于恢复分析。 |
| oneSentenceSummary | String | 否 | 一句话总结。 |
| keywords | JSON Array | 否 | 关键词列表。 |
| sentenceCount | Int | 否 | 句子数量。 |
| characterCount | Int | 否 | 字符数量。 |

历史记录存储在 SharedPreferences，名称为 `smartread_history`，键为 `records`。每次保存时，系统把新的记录放在列表首位，并过滤相同原文的旧记录，最多保留 12 条。清空历史记录时写入空数组。

## 附录：界面截图汇总

本文档已插入实际截图，不使用空白示意图。截图来源如下：

- 首页：`screenshots/app/V1.0首页_20260512.png`
- 文本输入：`screenshots/app/V1.0文本输入_20260512.png`
- 摘要结果：`screenshots/app/V1.0摘要结果_20260512.png`
- LiteRT 端侧分析：`screenshots/app/V1.0LiteRT端侧分析_20260512.png`
- Agent 问答：`screenshots/app/V1.0Agent问答_20260512.png`
- 知识卡片：`screenshots/app/V1.0知识卡片_20260512.png`
- 复习题：`screenshots/app/V1.0复习题_20260512.png`
- 历史记录：`screenshots/app/V1.0历史记录页_20260512.png`
"""


def set_doc_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[name].font.name = "黑体"


def add_para(doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    if align is not None:
        p.alignment = align
    return p


def add_heading_numbered(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        table.cell(0, i).text = value
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 5.6) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.runs:
        run.font.size = Pt(9)


def add_screenshot_grid(doc: Document, items: list[tuple[str, Path]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for i in range(0, len(items), 2):
        row = table.add_row().cells
        for j in range(2):
            if i + j >= len(items):
                continue
            title, path = items[i + j]
            para = row[j].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(title).bold = True
            para = row[j].add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(str(path), width=Inches(2.25))
    doc.add_paragraph()


def build_docx(diagrams: dict[str, Path]) -> None:
    doc = Document()
    set_doc_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    for text, size, bold in [
        ("SmartRead Agent 系统", 22, True),
        ("软件详细设计说明书", 26, True),
        ("Version 1.0", 14, False),
        ("文档名称：SmartReadAgent_软件详细设计说明书_模板版.docx", 12, False),
        ("课程：软件项目研发实践（3）", 12, False),
        ("指导老师：林立老师", 12, False),
        ("小组成员：林立洲（121072021030）、江轩宇（121052023075）", 12, False),
        ("提交说明：电子版与纸质版于 2026-07-10 前提交", 12, False),
    ]:
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
    doc.add_page_break()

    doc.add_heading("修订历史记录", level=1)
    add_table(
        doc,
        ["日期", "版本号", "修改说明", "修改人", "核准人"],
        [
            ["2026-05-11", "V0.1", "项目立项、目录和 Git/Gitee/GitHub 初始化", "林立洲", "林立洲"],
            ["2026-05-12", "V0.2", "完成文本输入、本地摘要、关键词和分点摘要 MVP", "林立洲", "林立洲"],
            ["2026-05-12", "V0.3", "增加 Agent 问答、知识卡片和复习题", "林立洲", "林立洲"],
            ["2026-05-12", "V0.4", "增加 LiteRT 端侧句子重要性分析模型", "林立洲", "林立洲"],
            ["2026-05-12", "V0.5", "增加本地历史记录保存、恢复和清空", "林立洲", "林立洲"],
            ["2026-05-13", "V1.0", "收口期末演示稳定版、APK、截图、测试和说明材料", "林立洲、江轩宇", "林立洲"],
            ["2026-06-29", "V1.0 初稿", "按课程模板整理软件详细设计说明书", "林立洲", "林立洲"],
        ],
    )

    doc.add_heading("目录", level=1)
    for line in [
        "1 引言",
        "1.1 目的与范围",
        "1.2 预期的读者",
        "1.3 系统的范围",
        "1.4 参考资料",
        "1.5 术语、缩写词",
        "1.6 模块命名规则",
        "2 建议的系统",
        "2.1 建议系统概述",
        "2.2 功能性需求概述",
        "2.3 非功能性需求",
        "2.3.1 用户界面与人员因素",
        "2.3.2 错误处理与极端情况",
        "2.3.3 质量要求",
        "2.3.4 用例模型",
        "2.3.5 对象模型",
        "3 模块总汇",
        "3.1 模块总汇表",
        "3.2 模块关系图",
        "3.3 模块设计清单",
        "3.4 附录：数据存储结构设计",
    ]:
        add_para(doc, line)
    doc.add_page_break()

    # Main content generated from the markdown, but with actual diagrams/screenshots in-place.
    sections = design_markdown().split("## 1 引言", 1)[1]
    # Manual writing keeps the template order and avoids markdown table parsing glitches.
    add_heading_numbered(doc, "1 引言", 1)
    add_heading_numbered(doc, "1.1 目的与范围", 2)
    add_para(doc, "本文档是 SmartRead Agent 系统的详细设计说明书，按照课程提供的“软件详细设计说明书模板”结构编写。文档用于说明系统设计目标、功能范围、非功能需求、用例模型、对象模型、模块关系、模块设计清单和数据存储结构，为《软件项目研发实践（3）》期末演示后的电子版与纸质版提交提供依据。")
    add_para(doc, "SmartRead Agent 当前版本为 V1.0 期末演示稳定版，已生成课程演示用 debug APK：release/SmartReadAgent-v1.0-debug.apk。该 APK 用于课程验收、真机测试和现场演示，不作为应用商店正式商业发布版本。本文档只描述当前已实现的功能，不把 OCR、云端大模型、登录、服务器、多设备同步等后续设想写成已实现内容。")
    add_para(doc, "本文档覆盖 Android/Kotlin 客户端、Jetpack Compose 界面、本地文本分析、本地规则型阅读 Agent、LiteRT 端侧句子重要性分析、本地历史记录和演示部署流程。重点突出移动端机器学习在应用中的真实整合方式，即模型文件随 APK 集成并在 Android 端侧完成推理。")
    add_heading_numbered(doc, "1.2 预期的读者", 2)
    add_para(doc, "本文档的预期读者包括课程指导老师、期末演示评审人员、小组成员、后续查看仓库的同学，以及需要理解系统模块设计和代码结构的人。读者不需要提前阅读所有源代码，但应能通过本文档理解系统主要模块、数据对象和运行流程。对于评分人员，本文档也可作为核对项目是否真实实现 LiteRT、Agent、历史记录和移动端演示能力的依据。")
    add_heading_numbered(doc, "1.3 系统的范围", 2)
    add_para(doc, "SmartRead Agent 面向学生阅读教材、课堂资料、技术文章和项目说明文档的场景，提供移动端阅读辅助能力。系统当前范围包括：文本输入、示例文本、本地摘要、关键词提取、分点摘要、LiteRT 端侧句子重要性分析、Agent 问答、知识卡片、复习题、本地历史记录保存、历史记录恢复和清空历史记录。")
    add_para(doc, "系统当前不包括：用户注册登录、云端同步、服务器后台、OCR 图片识别、语音输入、多设备同步、商业发布版签名包和真实云端大模型 API。这些内容不在当前演示范围内，后续可以作为扩展方向。")
    add_heading_numbered(doc, "1.4 参考资料", 2)
    for item in ["《2023级软工期末考核说明.pdf》", "《2023级软工期末课程演示流程.pdf》", "《2023级期末演示评分表.xlsx》", "README.md、docs/V0.4_LiteRT端侧模型集成说明.md、docs/关键代码与实现方法说明.md", "Android / Kotlin / Jetpack Compose 相关开发资料", "TensorFlow Lite / LiteRT Android 端侧推理相关资料"]:
        add_para(doc, item, style="List Bullet")
    add_heading_numbered(doc, "1.5 术语、缩写词", 2)
    add_table(doc, ["术语", "说明"], [
        ["SmartRead Agent", "本项目名称，表示面向阅读场景的智能辅助 App。"],
        ["LiteRT / TensorFlow Lite", "移动端轻量模型推理框架。"],
        ["本地规则型阅读 Agent", "当前版本 Agent 实现方式，主要使用本地规则和文章分析结果。"],
        ["ArticleAnalysis", "文章分析结果数据对象。"],
        ["SentenceImportance", "句子重要性评分结果对象。"],
        ["APK", "Android 安装包。"],
        ["SharedPreferences", "Android 本地轻量键值存储。"],
    ])
    add_heading_numbered(doc, "1.6 模块命名规则", 2)
    add_para(doc, "系统代码采用 Kotlin 命名规范。数据对象使用名词型 PascalCase，例如 ArticleAnalysis、SentenceImportance、KnowledgeCard、HistoryRecord。功能对象或仓库对象使用职责名加后缀，例如 TextAnalyzer、SentenceImportanceAnalyzer、SentenceImportanceClassifier、LocalReadingAgent、HistoryRepository。Compose 页面函数使用页面含义命名，例如 AgentChatScreen、KnowledgeCardsScreen、HistoryScreen。资源文件使用能够说明功能的名称，例如 sentence_importance_model.tflite。")

    add_heading_numbered(doc, "2 建议的系统", 1)
    add_heading_numbered(doc, "2.1 建议系统概述", 2)
    add_para(doc, "SmartRead Agent 是一个 Android/Kotlin 移动端 AI 阅读辅助应用。用户可以在首页输入文本或点击示例文本，然后点击“开始智能分析”。系统会在本地生成一句话总结、关键词和分点摘要，同时调用 LiteRT 模型对句子重要性进行端侧评分。分析完成后，用户可以进入 Agent 问答页面围绕当前文章提问，也可以进入知识卡片页面查看复习卡片和复习题，还可以在历史记录页面恢复最近的分析结果。")
    add_para(doc, "系统设计强调“轻量、稳定、可解释、适合课堂演示”。当前版本的 Agent 问答基于当前文章分析结果和本地规则实现，不依赖外部大模型接口，因此在没有网络的情况下也能完成演示。LiteRT 模块是系统移动机器学习部分的核心，模型文件放入 Android assets，运行时由 SentenceImportanceClassifier 加载并推理。")
    add_heading_numbered(doc, "2.2 功能性需求概述", 2)
    for item in [
        "用户可以手动输入待分析文本。",
        "用户可以选择内置示例文本，保证课堂演示稳定。",
        "系统可以生成一句话总结、关键词和分点摘要。",
        "系统可以显示字符数、句子数等基础统计信息。",
        "系统可以调用 LiteRT 端侧模型，对文章句子进行重要性评分。",
        "系统可以在模型加载失败时使用规则评分兜底，避免演示中断。",
        "系统可以进入 Agent 问答页面，支持快捷问题和手动提问。",
        "系统可以生成知识卡片和复习题。",
        "系统可以保存最近 12 条历史分析记录。",
        "用户可以恢复历史记录，也可以清空历史记录。",
    ]:
        add_para(doc, item, style="List Number")
    add_heading_numbered(doc, "2.3 非功能性需求", 2)
    add_para(doc, "系统需要满足现场可演示、离线可用、响应及时、界面清晰、异常可控和材料可追溯等非功能需求。由于课程演示每组不超过 15 分钟，App 操作路径必须简洁，示例文本必须可直接触发完整功能链路。由于教室网络环境不一定稳定，核心功能不依赖网络。")
    add_heading_numbered(doc, "2.3.1 用户界面与人员因素", 3)
    add_para(doc, "界面采用 Jetpack Compose 实现，整体为移动端纵向滚动布局。首页突出项目名称、版本号、输入区、示例文本按钮和“开始智能分析”按钮。分析结果以卡片方式展示，包含摘要、关键词、LiteRT 评分和功能入口。Agent、知识卡片和历史记录作为独立页面，便于演示时逐个展示。")
    add_para(doc, "用户角色主要是学生用户和小组演示者。学生用户关注能否快速理解文章；演示者关注流程是否稳定、老师是否能看清楚功能和机器学习亮点。江轩宇在现场负责手机端实际操作，林立洲负责同步讲解和答辩。")
    add_screenshot_grid(doc, [
        ("首页", screenshot("V1.0首页_20260512.png")),
        ("文本输入", screenshot("V1.0文本输入_20260512.png")),
    ])
    add_heading_numbered(doc, "2.3.2 错误处理与极端情况", 3)
    add_para(doc, "系统对空文本、模型加载失败、历史记录 JSON 解析失败和无分析结果进入子页面等情况进行了处理。空文本分析时不生成结果，而是提示用户先输入内容；LiteRT 模型加载或推理失败时，系统使用规则评分兜底；历史记录解析失败时返回空列表；用户未完成分析就进入 Agent 或知识卡片页面时，页面显示友好提示。")
    add_para(doc, "这些处理的目的不是追求复杂异常体系，而是保证课程现场演示不断线。只要 App 能打开，用户就能通过示例文本稳定走完整流程。")
    add_heading_numbered(doc, "2.3.3 质量要求", 3)
    add_table(doc, ["主要质量属性", "详细要求"], [
        ["正确性", "摘要、关键词、LiteRT 分析、Agent 回答、知识卡片和历史记录应基于当前输入文本生成，不展示与当前文章无关的内容。"],
        ["健壮性", "空输入、模型加载失败、历史记录为空、历史 JSON 解析失败等情况不应导致 App 崩溃。"],
        ["兼容性", "APK 可在 Android 模拟器和安卓真机上安装运行；课程演示用 debug APK 不作为商业发布版。"],
    ])
    add_heading_numbered(doc, "2.3.4 用例模型", 3)
    add_table(doc, ["参与者", "说明"], [
        ["学生用户", "输入或选择文本，查看摘要、LiteRT 分析、Agent 回答、知识卡片和历史记录。"],
        ["小组演示者", "在课堂现场操作 App，按固定流程展示功能和机器学习整合效果。"],
        ["老师", "查看系统演示、提出问题、根据评分表评价应用创新性、功能丰富度和机器学习整合程度。"],
    ])
    add_table(doc, ["用例名称", "参与者", "简要说明"], [
        ["输入或选择文本", "学生用户、小组演示者", "用户输入待分析文本或点击示例文本。"],
        ["生成阅读分析结果", "学生用户、小组演示者", "系统生成摘要、关键词、分点摘要和基础统计。"],
        ["执行 LiteRT 端侧分析", "学生用户、小组演示者", "系统对句子进行重要性评分并显示来源。"],
        ["使用 Agent 问答", "学生用户、小组演示者", "用户围绕当前文章点击快捷问题或手动提问。"],
        ["查看知识卡片和复习题", "学生用户、小组演示者", "系统根据分析结果生成复习材料。"],
        ["查看和恢复历史记录", "学生用户、小组演示者", "用户查看最近分析并恢复其中一条。"],
    ])
    add_image(doc, diagrams["use_case"], "图 2-1 SmartRead Agent 顶层用例图", width=5.8)
    for item in [
        "输入或选择文本：用户在首页文本框输入内容，或点击示例文本按钮。系统更新输入框内容并等待分析。",
        "生成阅读分析结果：用户点击“开始智能分析”，系统调用 TextAnalyzer.analyze 生成 ArticleAnalysis。",
        "执行 LiteRT 端侧分析：系统提取句子特征并调用 SentenceImportanceClassifier，在结果页显示分数和来源。",
        "使用 Agent 问答：用户进入 Agent 页面，选择快捷问题或输入问题，系统调用 LocalReadingAgent.answerQuestion 返回回答。",
        "查看知识卡片和复习题：用户进入知识卡片页面，系统根据当前 ArticleAnalysis 生成 KnowledgeCard 和 QuizQuestion。",
        "查看和恢复历史记录：用户进入历史记录页面，点击某条记录后恢复原文并重新分析。",
    ]:
        add_para(doc, item, style="List Bullet")
    add_heading_numbered(doc, "2.3.5 对象模型", 3)
    add_para(doc, "系统的主数据对象为 ArticleAnalysis，它把原文、摘要、关键词、分点摘要、句子数、字符数、LiteRT 结果和模型状态集中保存。Agent、知识卡片、复习题和历史记录模块均围绕 ArticleAnalysis 工作，避免重复解析原文。")
    add_image(doc, diagrams["object"], "图 2-2 SmartRead Agent 对象模型图", width=5.8)

    add_heading_numbered(doc, "3．模块总汇", 1)
    add_heading_numbered(doc, "3.1 模块总汇表", 2)
    add_table(doc, ["模块名称", "类名称 / 对象名称", "主要职责"], [
        ["文本输入模块", "Compose 首页、SampleTexts", "处理手动输入、示例文本选择、清空和开始分析。"],
        ["本地摘要与关键词模块", "TextAnalyzer", "切分句子、提取关键词、生成一句话总结和分点摘要。"],
        ["LiteRT 端侧分析模块", "SentenceImportanceAnalyzer、SentenceImportanceClassifier", "提取 5 维特征，加载 .tflite 模型并输出句子重要性评分。"],
        ["Agent 问答模块", "LocalReadingAgent", "根据当前文章分析结果进行本地规则型问答。"],
        ["知识卡片与复习题模块", "generateKnowledgeCards、generateQuizQuestions", "生成复习卡片和简答题。"],
        ["历史记录模块", "HistoryRepository", "使用 SharedPreferences 保存、读取、恢复和清空历史记录。"],
        ["UI 展示模块", "SmartReadAgentApp、AgentChatScreen、KnowledgeCardsScreen、HistoryScreen", "组织首页、结果页、Agent 页面、知识卡片页面和历史页面。"],
    ])
    add_heading_numbered(doc, "3.2 模块关系图", 2)
    add_image(doc, diagrams["module"], "图 3-1 SmartRead Agent 模块关系图", width=5.8)
    add_para(doc, "系统主流程为：文本输入模块收集原文，交给 TextAnalyzer 得到 ArticleAnalysis；分析结果继续进入 LiteRT 端侧分析模块、Agent 问答模块、知识卡片模块和历史记录模块；最终由 Compose UI 展示给用户。LiteRT 模型异常时，系统使用规则评分兜底，不影响摘要和其他功能。")

    add_heading_numbered(doc, "3.3 模块设计清单", 2)
    module_texts = [
        ("模块一：文本输入模块", "提供文本输入、示例文本选择、清空和开始分析入口。输入为用户手动输入或内置示例文本；输出为传给 TextAnalyzer 的原始字符串。异常处理包括空文本提示和清空后状态重置。", [("首页", screenshot("V1.0首页_20260512.png")), ("文本输入", screenshot("V1.0文本输入_20260512.png"))]),
        ("模块二：本地摘要与关键词模块", "生成一句话总结、关键词、分点摘要、字符数和句子数。核心对象为 TextAnalyzer。它先切分句子，再按关键词命中、句子长度和位置给句子评分。关键词提取结合领域词、英文技术词和中文短语窗口。", [("摘要结果", screenshot("V1.0摘要结果_20260512.png")), ("操作入口", screenshot("V1.0操作入口_20260512.png"))]),
        ("模块三：LiteRT 端侧句子重要性分析模块", "体现移动端机器学习整合。核心对象为 SentenceImportanceAnalyzer 和 SentenceImportanceClassifier。输入为句子、关键词、句子位置和总句数，输出为 SentenceImportance。模型输入 5 维特征：句子长度归一化、关键词重合度、位置分数、标点提示分数、总结提示词分数。模型文件为 android/app/src/main/assets/sentence_importance_model.tflite。", [("LiteRT 端侧分析", screenshot("V1.0LiteRT端侧分析_20260512.png"))]),
        ("模块四：Agent 问答模块", "围绕当前文章分析结果回答用户问题。当前版本统一说明为本地规则型阅读 Agent，回答内容主要来自摘要、关键词、复习题和知识卡片等本地分析结果。核心对象为 LocalReadingAgent。它根据问题意图识别返回概括、关键词、复习重点、复习题、知识卡片提示或综合回答。", [("Agent 问答", screenshot("V1.0Agent问答_20260512.png")), ("Agent 快捷问题", screenshot("V1.0Agent快捷问题_20260512.png"))]),
        ("模块五：知识卡片与复习题模块", "把分析结果转换为复习材料。输入为 ArticleAnalysis，输出为 KnowledgeCard 和 QuizQuestion。知识卡片包括概念、观点和复习类内容；复习题以简答题和参考答案形式展示。", [("知识卡片", screenshot("V1.0知识卡片_20260512.png")), ("复习题", screenshot("V1.0复习题_20260512.png"))]),
        ("模块六：历史记录模块", "保存最近 12 条分析记录，并支持恢复和清空。核心对象为 HistoryRepository。存储方式为 SharedPreferences 中的 JSON 数组，不使用数据库。该方案适合当前课程演示规模，后续如扩展为正式产品可迁移至 Room。", [("历史记录", screenshot("V1.0历史记录页_20260512.png")), ("历史恢复", screenshot("V1.0历史记录恢复_20260512.png"))]),
        ("模块七：UI 展示模块", "使用 Jetpack Compose 组织首页、结果页、Agent 页面、知识卡片页面和历史记录页面。页面切换由 Screen 枚举控制，UI 状态由 Compose 的 remember 和 mutableStateOf 管理。界面强调课程演示时的可读性和操作路径清晰。", [("清空历史记录", screenshot("V1.0清空历史记录_20260512.png"))]),
    ]
    for title, body, images in module_texts:
        add_heading_numbered(doc, title, 3)
        add_para(doc, body)
        if len(images) == 1:
            add_image(doc, images[0][1], f"图：{images[0][0]}", width=2.8)
        else:
            add_screenshot_grid(doc, images)

    add_heading_numbered(doc, "3.4 附录：数据存储结构设计", 2)
    add_table(doc, ["列名", "数据类型", "是否为空", "说明"], [
        ["id", "Long", "否", "历史记录唯一标识，使用保存时间戳。"],
        ["title", "String", "否", "历史记录标题，默认取一句话总结前 28 个字符。"],
        ["preview", "String", "否", "原文预览，列表页展示使用。"],
        ["savedAt", "Long", "否", "保存时间戳。"],
        ["displayTime", "String", "否", "格式化后的展示时间。"],
        ["originalText", "String", "否", "原始文本，用于恢复分析。"],
        ["oneSentenceSummary", "String", "否", "一句话总结。"],
        ["keywords", "JSON Array", "否", "关键词列表。"],
        ["sentenceCount", "Int", "否", "句子数量。"],
        ["characterCount", "Int", "否", "字符数量。"],
    ])
    add_para(doc, "历史记录存储在 SharedPreferences，名称为 smartread_history，键为 records。每次保存时，系统把新的记录放在列表首位，并过滤相同原文的旧记录，最多保留 12 条。清空历史记录时写入空数组。")

    locked = []
    for target in (ALT_DOCX, OUT_DOCX):
        try:
            doc.save(target)
        except PermissionError:
            locked.append(target)
    if locked:
        names = ", ".join(str(path) for path in locked)
        raise PermissionError(f"close the Word/WPS document before rebuilding: {names}")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    diagrams = make_diagrams()
    OUT_MD.write_text(design_markdown().strip() + "\n", encoding="utf-8")
    build_docx(diagrams)
    print(OUT_MD)
    print(OUT_DOCX)
    if ALT_DOCX.exists():
        print(ALT_DOCX)
    for path in diagrams.values():
        print(path)


if __name__ == "__main__":
    main()
