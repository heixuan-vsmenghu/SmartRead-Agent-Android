# -*- coding: utf-8 -*-
"""Export high-value Markdown notes to phone-friendly DOCX files."""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT / "docs" / "手机查看DOCX"

TARGETS = [
    ("01_期末汇报安排.docx", PROJECT / "docs" / "期末汇报安排.md", "期末汇报安排"),
    ("02_PPT逐页内容.docx", PROJECT / "docs" / "期末演示PPT逐页内容.md", "PPT 逐页内容"),
    ("03_期末汇报讲稿.docx", PROJECT / "docs" / "期末现场演示讲稿.md", "期末汇报讲稿"),
    ("04_问答参考.docx", PROJECT / "docs" / "答辩可能问题.md", "问答参考"),
    ("05_汇报前检查.docx", PROJECT / "docs" / "演示当天检查清单.md", "汇报前检查"),
    ("06_评分项对照表.docx", PROJECT / "docs" / "期末评分项对应自检表.md", "评分项对照表"),
    ("07_最终提交说明.docx", PROJECT / "docs" / "最终提交说明.md", "最终提交说明"),
    ("08_软件实践3_期末要求对照检查表.docx", PROJECT / "docs" / "软件实践3_期末要求对照检查表.md", "软件实践 3 期末要求对照检查表"),
    ("09_江轩宇手机录屏演示流程.docx", PROJECT / "demo" / "江轩宇手机录屏演示流程.md", "江轩宇手机录屏演示流程"),
    ("10_Android手机安装说明.docx", PROJECT / "release" / "Android手机安装说明.md", "Android 手机安装说明"),
    ("11_APK信息说明.docx", PROJECT / "release" / "APK信息说明.md", "APK 信息说明"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(clean_inline(text))
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.font.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(clean_inline(text))
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def is_table_block(lines: list[str], i: int) -> bool:
    if i + 1 >= len(lines):
        return False
    return lines[i].lstrip().startswith("|") and lines[i + 1].lstrip().startswith("|") and "---" in lines[i + 1]


def parse_table(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row):
            rows.append(row)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def convert_md_to_docx(source: Path, target: Path, title: str) -> None:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_doc(doc, title)
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                add_para(doc, "\n".join(code_lines), style=None)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            doc.add_paragraph("—" * 24)
            i += 1
            continue
        if is_table_block(lines, i):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            doc.add_heading(clean_inline(heading_match.group(2)), level=level)
            i += 1
            continue
        quote_match = re.match(r"^>\s*(.+)$", stripped)
        if quote_match:
            add_para(doc, "提示：" + quote_match.group(1))
            i += 1
            continue
        list_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if list_match:
            add_para(doc, "• " + list_match.group(1))
            i += 1
            continue
        task_match = re.match(r"^[-*]\s+\[[ xX]\]\s+(.+)$", stripped)
        if task_match:
            add_para(doc, "□ " + task_match.group(1))
            i += 1
            continue
        numbered_match = re.match(r"^(\d+)[.)]\s+(.+)$", stripped)
        if numbered_match:
            add_para(doc, f"{numbered_match.group(1)}. {numbered_match.group(2)}")
            i += 1
            continue
        add_para(doc, stripped)
        i += 1
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


def build_index() -> None:
    doc = Document()
    configure_doc(doc, "手机查看 DOCX 文件索引")
    add_para(doc, "这些文件由项目中的 Markdown 材料导出，适合在手机或微信文件里快速查看。建议汇报前优先看 01、02、03、04、05。")
    rows = [["序号", "文件", "用途"]]
    purposes = [
        "汇报顺序、时间表和分工",
        "PPT 每页内容",
        "PPT 逐页讲稿",
        "常见问题与回答参考",
        "汇报前材料检查",
        "课程评分项对照",
        "提交材料范围说明",
        "课程要求对照",
        "江轩宇手机视频操作步骤",
        "APK 安装步骤",
        "APK 文件信息与用途说明",
    ]
    for idx, (target_name, _source, title) in enumerate(TARGETS, 1):
        rows.append([f"{idx:02d}", target_name, purposes[idx - 1]])
    add_table(doc, rows)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DIR / "00_手机查看文件索引.docx")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_index()
    for target_name, source, title in TARGETS:
        if not source.exists():
            raise FileNotFoundError(source)
        convert_md_to_docx(source, OUT_DIR / target_name, title)
    print(OUT_DIR)
    for path in sorted(OUT_DIR.glob("*.docx")):
        print(path)


if __name__ == "__main__":
    main()
